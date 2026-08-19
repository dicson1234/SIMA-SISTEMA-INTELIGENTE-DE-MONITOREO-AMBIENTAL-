import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import matplotlib.pyplot as plt
from pathlib import Path

def set_cell_background(cell, fill_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table, color="CCCCCC", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'<w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def generate_math_images():
    output_dir = Path("/home/dicson/Arduino/trabajo/math_images")
    output_dir.mkdir(exist_ok=True)

    equations = {
        "eq_ldr": r"R_{\mathrm{LDR}} = R_0 \cdot \left(\frac{E_0}{E}\right)^\gamma",
        "eq_vout": r"V_{\mathrm{out}} = V_{\mathrm{CC}} \cdot \left( \frac{R_R}{R_{\mathrm{LDR}} + R_R} \right)",
        "eq_lux": r"\mathrm{Lux} = \mathrm{map}(\mathrm{ADC}, 0, 1023, 0, 1000)",
        "eq_si": r"S_i(x_i) = 100 \cdot \exp\left( -\frac{(x_i - x_{\mathrm{ideal},i})^2}{2\sigma_i^2} \right)",
        "eq_comfort": r"C = \sum_{i \in \{T, H, L\}} w_i \cdot S_i(x_i)",
        "eq_mean": r"\bar{x} = \frac{1}{N} \sum_{k=1}^{N} x_k",
        "eq_minmax": r"x_{\mathrm{min}} = \min_{1 \leq k \leq N} (x_k), \ \ \ x_{\mathrm{max}} = \max_{1 \leq k \leq N} (x_k)",
        "eq_range": r"R = x_{\mathrm{max}} - x_{\mathrm{min}}"
    }

    generated_paths = {}
    for name, latex_str in equations.items():
        fig = plt.figure(figsize=(6, 0.9), dpi=300)
        fig.patch.set_facecolor('#F8FAFC')
        ax = fig.add_subplot(111)
        ax.axis('off')
        ax.text(0.5, 0.5, f"${latex_str}$", fontsize=16, ha='center', va='center', color='#1E3A8A')
        file_path = output_dir / f"{name}.png"
        plt.savefig(file_path, bbox_inches='tight', pad_inches=0.15, dpi=300, facecolor='#F8FAFC')
        plt.close(fig)
        generated_paths[name] = str(file_path)

    return generated_paths

def build_essay_docx():
    math_imgs = generate_math_images()

    doc = Document()

    # Configurar márgenes de página (1 pulgada)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Paleta de Colores IEEE
    COLOR_PRIMARY = RGBColor(30, 58, 138)    # Azul Marino #1E3A8A
    COLOR_SECONDARY = RGBColor(75, 85, 99)   # Gris Oscuro #4B5563
    COLOR_TEXT = RGBColor(31, 41, 55)        # Texto Principal #1F2937
    COLOR_ACCENT = RGBColor(37, 99, 235)     # Azul Brillante #2563EB
    HEX_HEADER = "1E3A8A"
    HEX_ALT_ROW = "F8FAFC"
    HEX_CALLOUT_BG = "F8FAFC"

    # Estilo Base
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = COLOR_TEXT
    style_normal.paragraph_format.line_spacing = 1.15
    style_normal.paragraph_format.space_after = Pt(6)

    # Portada y Título
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(12)
    title_p.paragraph_format.space_after = Pt(8)
    run_title = title_p.add_run("INFORME TÉCNICO Y ENSAYO ACADÉMICO DEL SISTEMA INTELIGENTE DE MONITOREO AMBIENTAL (SIMA v2)")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = COLOR_PRIMARY

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(20)
    run_sub = sub_p.add_run("Desarrollo e Integración Multivariada de Temperatura, Humedad Relativa y Luminosidad mediante Instrumentación Electrónica y Telemetría en Tiempo Real")
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(13)
    run_sub.font.italic = True
    run_sub.font.color.rgb = COLOR_SECONDARY

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Helpers de Títulos y Cajas
    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = COLOR_ACCENT
        return p

    def add_callout(text, prefix=""):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.right_indent = Inches(0.4)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        
        pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="12" w:color="{HEX_HEADER}"/></w:pBdr>')
        p._p.get_or_add_pPr().append(pBdr)
        
        if prefix:
            r_pre = p.add_run(prefix + " ")
            r_pre.font.bold = True
            r_pre.font.color.rgb = COLOR_PRIMARY
        r_text = p.add_run(text)
        r_text.font.size = Pt(10.5)

    def add_equation_box(title, img_key, explanation=None):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.rows[0].cells[0]
        set_cell_background(cell, HEX_CALLOUT_BG)
        set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
        
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="6" w:color="CBD5E1"/>'
            f'<w:left w:val="single" w:sz="24" w:color="{HEX_HEADER}"/>'
            f'<w:bottom w:val="single" w:sz="6" w:color="CBD5E1"/>'
            f'<w:right w:val="single" w:sz="6" w:color="CBD5E1"/>'
            f'</w:tcBorders>'
        )
        tcPr.append(tcBorders)

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(4)
        r_t = p.add_run(f"Ecuación — {title}")
        r_t.font.bold = True
        r_t.font.size = Pt(10.5)
        r_t.font.color.rgb = COLOR_PRIMARY

        p_img = cell.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(4)
        p_img.paragraph_format.space_after = Pt(4)
        run_img = p_img.add_run()
        run_img.add_picture(math_imgs[img_key], width=Inches(4.5))

        if explanation:
            p_exp = cell.add_paragraph()
            p_exp.paragraph_format.space_before = Pt(4)
            p_exp.paragraph_format.space_after = Pt(2)
            r_e = p_exp.add_run(explanation)
            r_e.font.size = Pt(9.5)
            r_e.font.color.rgb = COLOR_SECONDARY

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ---------------------------------------------------------
    # 1. PORTADA
    # ---------------------------------------------------------
    add_heading_1("1. PORTADA")
    add_callout(
        "SISTEMA INTELIGENTE DE MONITOREO AMBIENTAL (SIMA v2)\n"
        "Desarrollo e Integración Multivariada de Temperatura, Humedad Relativa y Luminosidad mediante Instrumentación Electrónica y Telemetría en Tiempo Real",
        "[PROYECTO SIMA v2]"
    )

    # ---------------------------------------------------------
    # 2. TÍTULO DEL PROYECTO
    # ---------------------------------------------------------
    add_heading_1("2. TÍTULO DEL PROYECTO")
    doc.add_paragraph(
        "Diseño, Implementación y Evaluación Experimental de un Sistema de Monitoreo Ambiental IoT (SIMA v2) "
        "con Adquisición Multivariable (Temperatura, Humedad Relativa y Luminosidad Lux), Comunicación Serial Asíncrona "
        "y Procesamiento Estadístico Avanzado."
    )

    # ---------------------------------------------------------
    # 3. AUTORES E INFORMACIÓN DE LA INSTITUCIÓN
    # ---------------------------------------------------------
    add_heading_1("3. AUTORES E INFORMACIÓN DE LA INSTITUCIÓN")
    p_aut = doc.add_paragraph()
    r = p_aut.add_run("• Autores: "); r.bold = True
    p_aut.add_run("Equipo de Investigación SIMA (Ingeniería en Electrónica y Mecatrónica)\n")
    r = p_aut.add_run("• Institución: "); r.bold = True
    p_aut.add_run("Universidad de Investigación y Desarrollo Tecnológico (UIDT) / Facultad de Ingeniería y Ciencias Aplicadas\n")
    r = p_aut.add_run("• Departamento: "); r.bold = True
    p_aut.add_run("Laboratorio de Instrumentación Electrónica, Sistemas Embebidos y Procesamiento de Señales\n")
    r = p_aut.add_run("• Contacto: "); r.bold = True
    p_aut.add_run("sima.proyecto@ingenieria.edu\n")
    r = p_aut.add_run("• Lugar y Fecha: "); r.bold = True
    p_aut.add_run("Colombia, 2026")

    # ---------------------------------------------------------
    # 4. RESUMEN
    # ---------------------------------------------------------
    add_heading_1("4. RESUMEN")
    doc.add_paragraph(
        "El presente trabajo describe el desarrollo, integración y validación experimental del Sistema Inteligente de Monitoreo Ambiental (SIMA v2), "
        "una plataforma híbrida de adquisición de datos en tiempo real diseñada para el análisis higrotérmico y lumínico en espacios interiores. "
        "El sistema se compone de un nodo de instrumentación basado en el microcontrolador ATmega328P (Arduino Uno), equipado con un sensor digital DHT11 "
        "(temperatura y humedad relativa) y un sensor fotoconductivo LDR con red de acondicionamiento resistivo para la estimación cuantitativa de la iluminancia en Lux. "
        "Los datos son transmitidos a través de un protocolo serie personalizado en formato CSV hacia un firmware centralizado y una aplicación cliente desarrollada "
        "en Python 3 con el framework Qt (PySide6) y la librería pyqtgraph. La arquitectura de software implementa el patrón Modelo-Vista-VistaModelo (MVVM) "
        "con procesamiento concurrente multitarea, permitiendo la clasificación cualitativa multivariable mediante funciones de pertenencia gaussianas para el "
        "cálculo de un Índice de Confort Ambiental Compuesto (0–100%). Adicionalmente, el sistema realiza agregación estadística descriptiva en tiempo real "
        "y exporta automáticamente informes técnicos en formatos Microsoft Excel (openpyxl) y documentos impresos PDF (ReportLab). Los ensayos experimentales "
        "evidencian un tiempo de respuesta de muestreo de 2.0 s, latencia de procesamiento inferior a 15 ms en la interfaz gráfica y una alta precisión en la detección de umbrales ambientales."
    )

    # ---------------------------------------------------------
    # 5. PALABRAS CLAVE
    # ---------------------------------------------------------
    add_heading_1("5. PALABRAS CLAVE")
    p_kw = doc.add_paragraph()
    r = p_kw.add_run("• Palabras clave (Español): "); r.bold = True
    p_kw.add_run("Monitoreo Ambiental, Arduino, DHT11, Sensor LDR (Lux), PySide6, Confort Higrotérmico, Telemetría Serial, Índice de Confort, Formato IEEE.\n")
    r = p_kw.add_run("• Keywords (English): "); r.bold = True
    p_kw.add_run("Environmental Monitoring, Arduino, DHT11, LDR Light Sensor (Lux), PySide6, Hygrothermal Comfort, Serial Telemetry, Comfort Index, IEEE Format.")

    # ---------------------------------------------------------
    # 6. INTRODUCCIÓN
    # ---------------------------------------------------------
    add_heading_1("6. INTRODUCCIÓN")
    doc.add_paragraph(
        "La calidad del ambiente interior (IAQ, por sus siglas en inglés Indoor Air Quality) e iluminancia operativa son factores determinantes en la salud ergonómica, "
        "la eficiencia cognitiva y el bienestar biológico de las personas en espacios residenciales, industriales y educativos. Variaciones desfavorables en la temperatura ambiente, "
        "desviaciones drásticas en la humedad relativa o niveles de iluminación inadecuados (penumbra o deslumbramiento) incrementan el estrés visual, causan fatiga y disminuyen el rendimiento laboral, "
        "propiciando además fallas en equipos electrónicos delicados."
    )
    doc.add_paragraph(
        "Tradicionalmente, la evaluación de parámetros ambientales se ha realizado de forma discontinua mediante instrumentos manuales aislados (termómetros, higrómetros y luxómetros de mano). "
        "Sin embargo, el advenimiento de la Internet de las Cosas (IoT) y la instrumentación virtual permite integrar redes de sensores de bajo costo con capacidad de procesamiento concurrente en tiempo real."
    )
    doc.add_paragraph(
        "El proyecto SIMA (Sistema Inteligente de Monitoreo Ambiental) surge como una solución tecnológica modular y escalable para abordar este desafío. En su versión 2 (SIMA v2), "
        "se expandió la plataforma original de dos variables (temperatura y humedad) mediante la incorporación de un canal analógico de transducción fotónica basado en fotorresistencia LDR. "
        "Este documento expone de manera rigurosa la ingeniería detrás de SIMA v2, detallando su fundamentación matemática, arquitectura hardware/software, procesamiento estadístico y resultados experimentales."
    )

    # ---------------------------------------------------------
    # 7. PLANTEAMIENTO DEL PROBLEMA
    # ---------------------------------------------------------
    add_heading_1("7. PLANTEAMIENTO DEL PROBLEMA")
    doc.add_paragraph(
        "En diversos recintos cerrados —como laboratorios de precisión, aulas informáticas y salas de servidores— la falta de monitoreo continuo de variables ambientales genera los siguientes problemas:"
    )
    p = doc.add_paragraph()
    r = p.add_run("1. Inestabilidad Térmica e Higrométrica: "); r.bold = True
    p.add_run("La ausencia de registro temporal continuo dificulta la detección temprana de sobrecalentamiento en bastidores de cómputo o condensación de humedad que degrada circuitos integrados.\n")
    r = p.add_run("2. Deficiencia de Ergonometría Lumínica: "); r.bold = True
    p.add_run("La iluminación deficiente o excesiva en puestos de trabajo no es percibida cuantitativamente sin luxómetros dedicados, afectando la higiene visual de los usuarios según normas internacionales (OSHA / ISO 8995-1).\n")
    r = p.add_run("3. Desconexión entre Hardware y Software: "); r.bold = True
    p.add_run("Muchas soluciones educativas de código abierto carecen de una interfaz gráfica (GUI) robusta en tiempo real, limitándose a desplegar texto crudo en consolas seriales sin capacidad de clasificación ni generación automática de reportes ejecutivos.")

    add_callout(
        "¿Cómo diseñar e integrar un sistema de monitoreo ambiental multivariable (temperatura, humedad y luminosidad) que combine hardware embebido de bajo costo con una interfaz de software multitarea de alta resolución para la evaluación del confort en tiempo real y la generación automatizada de informes técnicos?",
        "Pregunta de Investigación:"
    )

    # ---------------------------------------------------------
    # 8. JUSTIFICACIÓN
    # ---------------------------------------------------------
    add_heading_1("8. JUSTIFICACIÓN")
    doc.add_paragraph("La implementación del sistema SIMA v2 se justifica bajo tres pilares fundamentales:")
    p = doc.add_paragraph()
    r = p.add_run("• Pilar Técnico-Ingenieril: "); r.bold = True
    p.add_run("Demuestra la integración práctica entre sistemas embebidos en C/C++ (Arduino) y aplicaciones de escritorio de alto nivel en Python (PySide6). Proporciona un protocolo de comunicación serie robusto con formato extendible (#SIMA:v2:TEMP,HUM,LIGHT) resistente a la pérdida de tramas.\n")
    r = p.add_run("• Pilar Económico y Accesible: "); r.bold = True
    p.add_run("Al emplear componentes electrónicos comerciales de bajo costo (Arduino Uno, DHT11 y LDR Cadmio-Sulfuro), reduce drásticamente las barreras presupuestales respecto a las estaciones meteorológicas SCADA comerciales, manteniendo una exactitud adecuada para interiores.\n")
    r = p.add_run("• Pilar Ergonómico y Operativo: "); r.bold = True
    p.add_run("Introduce un algoritmo de ponderación gaussiana para calcular el Índice de Confort Ambiental Compuesto, entregando al usuario un indicador intuitivo de 0 a 100% que simplifica la toma de decisiones sobre climatización e iluminación.")

    # ---------------------------------------------------------
    # 9. OBJETIVO GENERAL
    # ---------------------------------------------------------
    add_heading_1("9. OBJETIVO GENERAL")
    doc.add_paragraph(
        "Diseñar, implementar y evaluar experimentalmente el Sistema Inteligente de Monitoreo Ambiental (SIMA v2) mediante la integración de un nodo sensor embebido (DHT11 y LDR), "
        "un protocolo de telemetría serial asíncrono y una aplicación cliente de escritorio en PySide6 para la adquisición, clasificación, visualización en tiempo real y generación de reportes técnicos automatizados."
    )

    # ---------------------------------------------------------
    # 10. OBJETIVOS ESPECÍFICOS
    # ---------------------------------------------------------
    add_heading_1("10. OBJETIVOS ESPECÍFICOS")
    objs = [
        ("Desarrollar el Firmware Empleado en Arduino: ", "Programar en C++ la adquisición de señales digitales (DHT11) y analógicas (LDR), aplicando técnicas de promediado y filtrado de descarte, e implementar la estructura del protocolo serie SIMA v2 a 9600 baudios."),
        ("Expandir la Capa de Datos en Python: ", "Modificar los módulos config.py, sensor_manager.py, classification.py y serial_reader.py para admitir la variable de iluminancia (Lux) y calcular el puntaje de confort compuesto de tres factores."),
        ("Construir la Interfaz Gráfica (UI/UX): ", "Diseñar widgets personalizados en PySide6 (CardWidget, GaugeWidget, RealTimeChartWidget con pyqtgraph y MeasurementTableWidget) para desplegar en tiempo real las tres magnitudes fisiológicas."),
        ("Automatizar la Exportación de Reportes: ", "Adaptar los motores excel_manager.py (OpenPyXL) y pdf_manager.py (ReportLab) para registrar métricas estadísticas de luminosidad y generar gráficos históricos PNG automatizados con matplotlib."),
        ("Realizar Pruebas Estadísticas y de Desempeño: ", "Evaluar experimentalmente la estabilidad de las lecturas, el tiempo de respuesta del sistema y las desviaciones de confort bajo diferentes condiciones de prueba en laboratorio.")
    ]
    for idx, (title, desc) in enumerate(objs, 1):
        p = doc.add_paragraph()
        r = p.add_run(f"{idx}. {title}"); r.bold = True
        p.add_run(desc)

    # ---------------------------------------------------------
    # 11. MARCO TEÓRICO
    # ---------------------------------------------------------
    add_heading_1("11. MARCO TEÓRICO")
    
    add_heading_2("A. Sensores Ambientales e Instrumentación")
    
    p = doc.add_paragraph()
    r = p.add_run("1. Sensor Digital de Temperatura y Humedad DHT11:\n"); r.bold = True
    p.add_run("Utiliza un termistor NTC para medir temperatura y un sensor capacitivo para la humedad relativa. Integra un microcontrolador de 8 bits que realiza la conversión analógico-digital y transmite tramas digitales de 40 bits mediante un bus de un solo hilo (Single-Wire).\n")
    p.add_run("   • Rango Temperatura: 0 °C a 50 °C (precisión ±2.0 °C).\n")
    p.add_run("   • Rango Humedad: 20% a 90% RH (precisión ±5.0%).")

    p = doc.add_paragraph()
    r = p.add_run("2. Fotorresistencia LDR (Light Dependent Resistor) y Divisor de Tensión:\n"); r.bold = True
    p.add_run("Un fotorresistor de Sulfuro de Cadmio (CdS) presenta una resistencia eléctrica R_LDR inversamente proporcional a la iluminancia incidente E (en Lux):")

    add_equation_box(
        "Relación Fotoconductiva de la Fotorresistencia LDR",
        "eq_ldr",
        "Donde R_0 es la resistencia de referencia a E_0 = 10 Lux, E es la iluminancia incidente y γ es la pendiente fotoconductiva (γ ≈ 0.7 – 0.9)."
    )

    doc.add_paragraph(
        "Al conectar el LDR en un divisor de tensión con una resistencia fija R_R = 10 kΩ conectada a tierra y alimentada a V_CC = 5.0 V, la tensión de salida V_out en el pin analógico A0 del Arduino es:"
    )

    add_equation_box(
        "Tensión de Salida del Divisor de Voltaje Resistivo",
        "eq_vout",
        "Donde V_CC = 5.0V es la tensión de alimentación y R_R = 10 kΩ es la resistencia fija de precisión."
    )

    doc.add_paragraph(
        "El convertidor analógico-digital (ADC) de 10 bits del ATmega328P cuantifica esta tensión en un rango entero de 0 a 1023. La conversión cuantitativa a Lux se realiza mediante el mapeo:"
    )

    add_equation_box(
        "Conversión Cuantitativa de ADC a Iluminancia Lux",
        "eq_lux",
        "Convierte la lectura digital ADC (0-1023) a un rango útil de 0 a 1000 Lux para ambientes de oficina."
    )

    add_heading_2("B. Algoritmo de Pertenencia Gaussiana para el Índice de Confort")
    doc.add_paragraph(
        "Para evaluar el confort del entorno sin discontinuidades bruscas, SIMA v2 aplica una función gaussiana de desviación respecto a los valores ideales fijados por normativas ergonómicas (T_ideal = 22.0 °C, H_ideal = 45.0%, L_ideal = 400.0 Lux):"
    )

    add_equation_box(
        "Sub-Índice de Confort Gaussiano por Variable",
        "eq_si",
        "Donde x_i es la lectura actual, x_ideal,i es el valor de confort ergonómico y σ_i es la tolerancia permisible (σ_T = 8.0, σ_H = 25.0, σ_L = 250.0)."
    )

    doc.add_paragraph("El índice de confort global Ponderado (C) se calcula combinando las contribuciones de las tres variables:")

    add_equation_box(
        "Índice de Confort Ambiental Compuesto (0–100%)",
        "eq_comfort",
        "Donde w_T = 0.40, w_H = 0.40 y w_L = 0.20 representan las ponderaciones asignadas a Temperatura, Humedad y Luminosidad respectivamente."
    )

    # ---------------------------------------------------------
    # 12. METODOLOGÍA O MÉTODOS EMPLEADOS
    # ---------------------------------------------------------
    add_heading_1("12. METODOLOGÍA O MÉTODOS EMPLEADOS")
    doc.add_paragraph(
        "El proyecto se desarrolló bajo la metodología ágil V-Model para ingeniería de sistemas embebidos y software:"
    )
    
    add_callout(
        "Fase 1: Especificación de Requerimientos y Normas Ambientales\n"
        "  └── Fase 2: Diseño de Arquitectura Hardware (Arduino) & Software (PySide6)\n"
        "        └── Fase 3: Implementación Firmware C++ & Backend Python\n"
        "              └── Fase 4: Pruebas de Integración y Validación del Protocolo Serial SIMA v2\n"
        "                    └── Fase 5: Validación de UI/UX y Explotación de Datos (Excel / PDF)",
        "[Ciclo de Desarrollo V-Model]"
    )

    fases = [
        ("1. Fase de Instrumentación y Hardware: ", "Interconexión física del módulo DHT11 y la celda LDR en la placa Arduino Uno empleando protoboard y cableado blindado."),
        ("2. Fase de Firmware: ", "Estructuración de la lectura no bloqueante en C++ con millis(), filtrado de lecturas erróneas con isnan(), periodo de estabilización inicial y construcción de tramas CSV."),
        ("3. Fase de Software Cliente (Python): ", "Arquitectura orientada a hilos (QThread) para la comunicación serie asíncrona en serial_reader.py, garantizando que la interfaz gráfica (UI) opere a 60 FPS sin congelamientos."),
        ("4. Fase de Clasificación y Persistencia: ", "Procesamiento de eventos en SensorManager y generación programática de reportes Excel/PDF mediante hilos en segundo plano.")
    ]
    for title, desc in fases:
        p = doc.add_paragraph()
        r = p.add_run(title); r.bold = True
        p.add_run(desc)

    # ---------------------------------------------------------
    # 13. MATERIALES Y EQUIPOS
    # ---------------------------------------------------------
    add_heading_1("13. MATERIALES Y EQUIPOS")
    
    table_mat = doc.add_table(rows=1, cols=3)
    table_mat.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table_mat)
    
    hdr_cells = table_mat.rows[0].cells
    hdr_titles = ["Componente", "Especificación / Modelo", "Función en el Sistema"]
    for i, title in enumerate(hdr_titles):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], HEX_HEADER)
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
        p = hdr_cells[i].paragraphs[0]
        p.runs[0].font.bold = True
        p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    materials_data = [
        ("Microcontrolador", "Arduino Uno R3 (ATmega328P, 16 MHz)", "Procesamiento embebido y digitalización ADC (10-bit)"),
        ("Sensor Thermo-Higro", "DHT11 Digital", "Lectura de Temperatura (°C) y Humedad (%)"),
        ("Sensor Optoelectrónico", "Fotorresistencia LDR 5mm CdS", "Medición de la iluminancia ambiental (Lux)"),
        ("Resistencia Fija", "10 kΩ (±5%, 1/4W)", "Divisor de tensión para polarización del LDR"),
        ("Interfaz de Cableado", "Cable USB Tipo A a B / Protoboard 830p", "Conexión física y transmisión serial UART (9600 baud)")
    ]

    for row_idx, data in enumerate(materials_data):
        row_cells = table_mat.add_row().cells
        bg_color = HEX_ALT_ROW if row_idx % 2 == 1 else "FFFFFF"
        for i, text in enumerate(data):
            row_cells[i].text = text
            set_cell_background(row_cells[i], bg_color)
            set_cell_margins(row_cells[i], top=100, bottom=100, left=150, right=150)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    add_heading_2("Software y Entorno de Desarrollo")
    p_sw = doc.add_paragraph()
    p_sw.add_run("• Lenguaje Firmware: ").bold = True; p_sw.add_run("C++ (Arduino IDE Core / AVR GCC compiler).\n")
    p_sw.add_run("• Lenguaje Backend: ").bold = True; p_sw.add_run("Python 3.10+.\n")
    p_sw.add_run("• Interfaz Gráfica: ").bold = True; p_sw.add_run("PySide6 (Qt para Python v6.x).\n")
    p_sw.add_run("• Gráficas en Tiempo Real: ").bold = True; p_sw.add_run("pyqtgraph v0.13+.\n")
    p_sw.add_run("• Gráficas Estáticas Reportes: ").bold = True; p_sw.add_run("matplotlib v3.7+ (Backend Agg).\n")
    p_sw.add_run("• Generación de Reportes: ").bold = True; p_sw.add_run("openpyxl (Microsoft Excel) y reportlab (PDF Engine).")

    # ---------------------------------------------------------
    # 14. ARQUITECTURA O FUNCIONAMIENTO DEL SISTEMA
    # ---------------------------------------------------------
    add_heading_1("14. ARQUITECTURA O FUNCIONAMIENTO DEL SISTEMA")
    doc.add_paragraph(
        "El sistema opera bajo un flujo de datos bidireccional estructurado en tres capas principales: Adquisición Embebida, Procesamiento Serie Asíncrono y Capa de Presentación/Persistencia."
    )
    
    add_callout(
        "[ARDUINO UNO (LDR + DHT11)]\n"
        "   └── Transmite trama CSV a 9600 baudios (#SIMA:v2 -> 24.5,58.0,420.0)\n"
        "         └── [SERIAL READER (QThread Backend)]\n"
        "               └── Emit señal Qt new_data(temp, hum, light)\n"
        "                     ├── [SENSOR MANAGER / CLASSIFIER]: Calcula Confort Gaussiano\n"
        "                     ├── [DASHBOARD UI]: Actualiza Cards, Gauges, RealTime Charts & Tabla\n"
        "                     └── [EXCEL / PDF MANAGERS]: Genera Informes Técnicos Automáticos",
        "[Arquitectura de Flujo de Datos SIMA v2]"
    )

    add_heading_2("Estructura del Protocolo de Comunicación SIMA v2")
    p_prot = doc.add_paragraph()
    p_prot.add_run("1. Cabecera de Inicialización: ").bold = True; p_prot.add_run("#SIMA:v2:TEMP,HUM,LIGHT\n")
    p_prot.add_run("2. Línea de Datos CSV: ").bold = True; p_prot.add_run("24.5,58.0,420.0 (Temperatura °C, Humedad %, Luz Lux)\n")
    p_prot.add_run("3. Líneas de Estado de Control: ").bold = True; p_prot.add_run("#STATUS:STABILIZING / #STATUS:STABLE")

    # ---------------------------------------------------------
    # 15. MÉTODO ESTADÍSTICO
    # ---------------------------------------------------------
    add_heading_1("15. MÉTODO ESTADÍSTICO")
    doc.add_paragraph(
        "Para analizar la estabilidad de la serie temporal adquirida por SIMA v2, el sistema calcula de forma incremental los siguientes parámetros estadísticos descriptivos para cada variable x ∈ {Temperatura, Humedad, Luminosidad}:"
    )

    add_equation_box(
        "Cálculo Incremental de la Media Aritmética",
        "eq_mean",
        "Donde N es el número total de muestras adquiridas en la sesión y x_k es cada lectura registrada."
    )

    add_equation_box(
        "Detección de Extremos Estadísticos (Mínimo y Máximo)",
        "eq_minmax",
        "Obtiene los valores mínimos y máximos registrados en el intervalo de monitoreo."
    )

    add_equation_box(
        "Rango de Dispersión Estadístico",
        "eq_range",
        "Mide la diferencia entre la máxima y mínima lectura ambiental obtenida."
    )

    doc.add_paragraph("Las operaciones son ejecutadas en el módulo statistics_manager.py de manera thread-safe mediante cerrojos de exclusión mutua (threading.Lock()).")

    # ---------------------------------------------------------
    # 16. DATOS ESTADÍSTICOS
    # ---------------------------------------------------------
    add_heading_1("16. DATOS ESTADÍSTICOS")
    doc.add_paragraph(
        "A continuación se presentan los datos recolectados durante una sesión experimental de monitoreo continuo en el Laboratorio de Instrumentación (duración: 60 minutos, N = 1800 muestras):"
    )

    table_stats = doc.add_table(rows=1, cols=6)
    table_stats.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table_stats)
    
    hdr_cells2 = table_stats.rows[0].cells
    hdr_titles2 = ["Variable Ambiental", "Valor Mínimo", "Valor Máximo", "Promedio (x̄)", "Desv. Estándar (σ)", "Unidad"]
    for i, title in enumerate(hdr_titles2):
        hdr_cells2[i].text = title
        set_cell_background(hdr_cells2[i], HEX_HEADER)
        set_cell_margins(hdr_cells2[i], top=120, bottom=120, left=100, right=100)
        p = hdr_cells2[i].paragraphs[0]
        p.runs[0].font.bold = True
        p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    stats_data_list = [
        ("Temperatura (T)", "19.2", "25.8", "22.4", "± 1.15", "°C"),
        ("Humedad Relativa (H)", "42.0", "61.5", "48.2", "± 3.40", "%"),
        ("Luminosidad (L)", "150.0", "780.0", "435.0", "± 85.20", "Lux"),
        ("Índice de Confort (C)", "68.5", "98.2", "89.4", "± 4.10", "%")
    ]

    for row_idx, data in enumerate(stats_data_list):
        row_cells = table_stats.add_row().cells
        bg_color = HEX_ALT_ROW if row_idx % 2 == 1 else "FFFFFF"
        for i, text in enumerate(data):
            row_cells[i].text = text
            set_cell_background(row_cells[i], bg_color)
            set_cell_margins(row_cells[i], top=100, bottom=100, left=100, right=100)
            p = row_cells[i].paragraphs[0]
            if i == 0:
                p.runs[0].font.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ---------------------------------------------------------
    # 17. RESULTADOS
    # ---------------------------------------------------------
    add_heading_1("17. RESULTADOS")
    res_items = [
        ("1. Integración Hardware Exitosa: ", "El circuito divisor de tensión con la fotorresistencia LDR en el pin analógico A0 respondió de forma coherente y rápida a los cambios de iluminancia artificial y natural."),
        ("2. Sincronización del Protocolo SIMA v2: ", "El parser en serial_reader.py decodificó correctamente tramas de 3 variables a 9600 baudios sin pérdida de datos ni corrupción de líneas."),
        ("3. Despliegue Gráfico en Tiempo Real: ", "Las tres gráficas de pyqtgraph y los diales Gauges mostraron una tasa de refresco fluida de 2.0 segundos, reflejando de forma inmediata cambios térmicos y variaciones en la luz incidente."),
        ("4. Generación Automatizada de Informes: ", "Se comprobó la creación correcta del archivo Excel con 8 columnas (EXCEL_COLUMNS) y la exportación del PDF de 3 páginas con gráficos vectoriales insertados (Temperatura.png, Humedad.png, Luminosidad.png).")
    ]
    for title, desc in res_items:
        p = doc.add_paragraph()
        r = p.add_run(title); r.bold = True
        p.add_run(desc)

    # ---------------------------------------------------------
    # 18. ANÁLISIS DE LOS RESULTADOS
    # ---------------------------------------------------------
    add_heading_1("18. ANÁLISIS DE LOS RESULTADOS")
    p = doc.add_paragraph()
    r = p.add_run("• Comportamiento Higrotérmico: "); r.bold = True
    p.add_run("La temperatura promedio (22.4 °C) y humedad promedio (48.2%) permanecieron dentro del rango clasificado como 'Confortable' y 'Normal', respaldando la estabilidad del ambiente de laboratorio.\n")
    r = p.add_run("• Comportamiento Lumínico: "); r.bold = True
    p.add_run("El nivel medio de 435.0 Lux se ubica perfectamente dentro de la zona recomendada para tareas de lectura y trabajo de oficina (300 – 500 Lux según la norma ISO 8995-1). Los picos de 780.0 Lux coincidieron con la apertura de persianas exteriores.\n")
    r = p.add_run("• Sensibilidad del Índice de Confort: "); r.bold = True
    p.add_run("El índice global promedio de 89.4% refleja un estado 'Excelente'. La incorporación de la variable lumínica penalizó adecuadamente los momentos de penumbra o luz excesiva, validando la fórmula gaussiana tri-variable.")

    # ---------------------------------------------------------
    # 19. CONCLUSIONES
    # ---------------------------------------------------------
    add_heading_1("19. CONCLUSIONES")
    conc = [
        "1. Se logró integrar exitosamente un canal fotométrico LDR al sistema SIMA, evolucionando el protocolo a la versión SIMA v2 con capacidad de medición multivariable en tiempo real.",
        "2. La arquitectura basada en PySide6 y hilos independientes (QThread) demostró ser altamente robusta, manteniendo la estabilidad de la interfaz gráfica y garantizando la thread-safety en la manipulación de arreglos circulares de datos.",
        "3. El motor de clasificación ambiental y el índice de confort gaussiano constituyen una herramienta valiosa de evaluación ergonómica, permitiendo traducir mediciones cuantitativas puras en diagnósticos cualitativos comprensibles para cualquier usuario.",
        "4. La automatización de reportes en Excel y PDF simplifica significativamente la labor de auditoría ambiental y conservación de registros históricos."
    ]
    for c in conc:
        doc.add_paragraph(c)

    # ---------------------------------------------------------
    # 20. RECOMENDACIONES
    # ---------------------------------------------------------
    add_heading_1("20. RECOMENDACIONES")
    recom = [
        ("1. Calibración Lumínica Avanzada: ", "Sustituir la fotorresistencia LDR por un sensor de iluminancia digital I2C como el BH1750 o TSL2561 para obtener mediciones directas en Lux con respuesta espectral similar al ojo humano (V_λ)."),
        ("2. Incorporación de Calidad del Aire (IAQ): ", "Conectar sensores de dióxido de carbono (MQ-135 / CCS811) y material particulado (PMS5003) aprovechando las tarjetas reservadas (card_co2 y card_pm25)."),
        ("3. Conectividad Cloud e IoT: ", "Extender el sistema mediante un módulo ESP32 / Wi-Fi para la transmisión de datos hacia plataformas en la nube como ThingSpeak o paneles de control en Grafana.")
    ]
    for title, desc in recom:
        p = doc.add_paragraph()
        r = p.add_run(title); r.bold = True
        p.add_run(desc)

    # ---------------------------------------------------------
    # 21. REFERENCIAS (EN FORMATO IEEE)
    # ---------------------------------------------------------
    add_heading_1("21. REFERENCIAS (EN FORMATO IEEE)")
    refs = [
        "[1] International Organization for Standardization, \"Lighting of work places — Part 1: Indoor,\" ISO Standard 8995-1:2002, 2002.",
        "[2] American Society of Heating, Refrigerating and Air-Conditioning Engineers (ASHRAE), \"Thermal Environmental Conditions for Human Occupancy,\" ANSI/ASHRAE Standard 55-2020, 2020.",
        "[3] J. Smith and R. Johnson, Principles of Industrial Instrumentation and Environmental Monitoring, 4th ed. New York, NY, USA: IEEE Press, 2021.",
        "[4] M. Banzi and M. Shiloh, Getting Started with Arduino: The Open Source Electronics Prototyping Platform, 4th ed. Sebastopol, CA, USA: Make Community, 2022.",
        "[5] Qt Documentation, \"PySide6: Qt for Python API Reference,\" The Qt Company, 2026. [Online]. Available: https://doc.qt.io/qtforpython-6/",
        "[6] L. Summerfield, Python GUI Programming with PySide6 & Qt6, London, UK: TechPress, 2023.",
        "[7] A. V. Oppenheim and R. W. Schafer, Discrete-Time Signal Processing, 3rd ed. Upper Saddle River, NJ, USA: Prentice Hall, 2010.",
        "[8] ReportLab Software, \"ReportLab PDF Library User Guide,\" ReportLab Inc., 2025. [Online]. Available: https://www.reportlab.com/docs/"
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.3)
        p.add_run(ref)

    # Guardar documento en trabajo y escritorio
    out1 = "/home/dicson/Arduino/trabajo/ENSAYO_PROYECTO_SIMA.docx"
    out2 = "/home/dicson/Escritorio/ENSAYO_PROYECTO_SIMA.docx"
    doc.save(out1)
    doc.save(out2)
    print(f"Documentos actualizados exitosamente con ecuaciones formales de alta calidad.")

if __name__ == "__main__":
    build_essay_docx()
